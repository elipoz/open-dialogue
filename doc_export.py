"""
Export dialogue to a Word document (.docx).
"""

import io

from docx import Document


def export_dialogue_to_docx(dialogue: list, speaker_labels: dict) -> bytes:
    """Build a Word document with the full conversation in chronological order.
    Format per message: <date-and-time> Speaker: Message.
    """
    doc = Document()
    # Ensure chronological order (oldest first)
    entries = list(dialogue)
    if entries and all(len(e) >= 3 and e[2] is not None for e in entries):
        entries = sorted(entries, key=lambda e: e[2])
    for entry in entries:
        party, content = entry[0], entry[1]
        ts = entry[2] if len(entry) >= 3 else None
        label = entry[3] if len(entry) >= 4 and entry[3] else speaker_labels.get(party, party)
        date_time = ts.strftime("%Y-%m-%d %H:%M") if ts else ""
        message = (content or "").rstrip()
        if message and not message.endswith("."):
            message += "."
        line = f"{date_time} {label}: {message}".strip()
        doc.add_paragraph(line)
        doc.add_paragraph()  # blank line between messages
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
